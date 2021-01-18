import docx  # pip install python-docx
import os
import pandas as pd


def get_data1(biaoges):
    """
    1.读取单个文件中规范的表格数据（第一个表格）
    :param biaoges:
    :return:
    """
    # 读取第一个（规范）表格，可能有多行数据
    rowi = len(biaoges[0].rows)
    print(rowi)

    # 定义空列表
    list1 = []
    for i in range(1, rowi):  # 从第2行开始循环
        list1.append([biaoges[0].cell(i, 0).text, biaoges[0].cell(i, 1).text, biaoges[0].cell(i, 2).text])

    print(list1)


def get_data2(biaoges):
    """
    2.读取单个文件中不规范的表格数据（第二个表格）
    :param biaoges:
    :return:
    """
    cells = biaoges[1]._cells
    cells_lis = [[cell.text for cell in cells]]
    print(cells_lis)

    """
        [['姓名', '郭达', '年龄', '22', '照片', '籍贯', '辽宁沈阳',
        '住址', '安徽合肥', '照片', '工作单位', 'XXX公司', '电话', '13333333333',
        '照片', '是否党员', '否', '出生日期', '1997-12-22', '照片']]
    """

    datai = pd.DataFrame(cells_lis)
    print('******************* datai完整 *******************')
    print(datai)
    print('******************* datai完整 *******************')

    datai = datai[[1, 3, 6, 8, 11, 13, 16, 18]]
    datai.columns = ['姓名', '年龄', '籍贯', '住址', '工作单位', '电话', '是否党员', '出生日期']
    print('******************* datai过滤 *******************')
    print(datai)
    print('******************* datai过滤 *******************')


def get_data_batch():
    # 1.3 批量读取数据
    new_dir = "D:" + os.sep + "sampple" + os.sep + "word_data" + os.sep
    os.chdir(new_dir)
    list2 = []
    for file in os.listdir('.'):
        if file.endswith('.docx'):
            doc = docx.Document('./' + file)
            biaoges = doc.tables
            rowi = len(biaoges[0].rows)
            for i in range(1, rowi):
                list2.append([biaoges[0].cell(i, 0).text, biaoges[0].cell(i, 1).text, biaoges[0].cell(i, 2).text])

    data1 = pd.DataFrame(list2, columns=['书籍名称', '种类', '价格'])
    print(data1)


if __name__ == '__main__':
    # 读取word文件
    filepath = "D:" + os.sep + "sampple" + os.sep + "word_data" + os.sep + "data.docx"
    doc = docx.Document(filepath)

    # 获取文档中所有表格对象的列表
    biaoges = doc.tables

    print('类型：', type(biaoges))  # 类型： <class 'list'>
    print('biaoges: ', biaoges)

    # biaoges: [<docx.table.Table object at 0x000001F827264DC8>, <docx.table.Table object at 0x000001F827286EC8>]

    # 处理规范的表格
    get_data1(biaoges)

    # 处理不规范的表格
    get_data2(biaoges)

    # 批量处理表格
    get_data_batch()
